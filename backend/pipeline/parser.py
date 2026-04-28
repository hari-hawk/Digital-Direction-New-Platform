"""Document parser — 3 paths: Docling (visual), raw text (mainframe), pandas (structured).

Produces ParsedDocument with sections ready for LLM extraction.
"""

import logging
import re
from pathlib import Path

import pdfplumber
import pandas as pd

from backend.config_loader import get_config_store, FormatConfig
from backend.models.schemas import ParsedDocument, ParsedSection

logger = logging.getLogger(__name__)


# Hard cap on the text size of a single section before it's sent to the LLM.
# Gemini Flash accepts ~1M input tokens; ~4 chars per token gives ~250K tokens
# per 1M chars. We cap well under that (~150K tokens) so prompt + global_context
# + few-shot examples + output budget all fit comfortably with headroom for
# carriers that have NO chunking config (Verizon, etc.) and would otherwise
# emit one giant "full_document" section that crashes Gemini.
#
# Sections are split on newline boundaries to avoid mid-row splits in tabular
# text. If a single line is itself >MAX_SECTION_CHARS the line is hard-split.
MAX_SECTION_CHARS = 600_000


# ============================================
# Main Router
# ============================================

def parse_document(
    file_path: str,
    carrier: str,
    document_type: str,
    format_variant: str | None = None,
) -> ParsedDocument:
    """Route to correct parser based on format config processing_path."""
    store = get_config_store()
    fmt_config = None

    if format_variant:
        # Find format config by name
        for fmt_name, fmt in store.get_formats(carrier).items():
            if fmt.name == format_variant:
                fmt_config = fmt
                break

    if not fmt_config:
        # Try matching format from file content
        try:
            text = _extract_raw_text(file_path, max_pages=2)
            fmt_config = store.match_format_variant(carrier, text)
        except Exception:
            pass

    # Determine processing path
    suffix = Path(file_path).suffix.lower()
    if suffix in (".csv", ".tsv", ".xlsx", ".xls"):
        processing_path = "pandas"
    elif suffix in (".msg", ".eml"):
        processing_path = "email"
    elif suffix == ".docx":
        processing_path = "docx"
    elif fmt_config:
        processing_path = fmt_config.processing_path
    else:
        processing_path = "raw_text"  # safe default for PDFs

    logger.info(f"Parsing {Path(file_path).name}: path={processing_path}, format={fmt_config.name if fmt_config else 'unknown'}")

    if processing_path == "docling":
        return parse_with_docling(file_path, carrier, document_type, fmt_config)
    elif processing_path == "raw_text":
        return parse_raw_text(file_path, carrier, document_type, fmt_config)
    elif processing_path == "pandas":
        return parse_structured_data(file_path, carrier, document_type, fmt_config)
    elif processing_path == "email":
        return parse_email(file_path, carrier, document_type)
    elif processing_path == "docx":
        return parse_docx(file_path, carrier, document_type)
    else:
        logger.warning(f"Unknown processing path: {processing_path}. Falling back to raw_text.")
        return parse_raw_text(file_path, carrier, document_type, fmt_config)


# ============================================
# Path 1: Docling (Visual Documents)
# ============================================

def parse_with_docling(
    file_path: str,
    carrier: str,
    document_type: str,
    fmt_config: FormatConfig | None = None,
) -> ParsedDocument:
    """Parse visual documents (invoices, contracts) with Docling layout analysis."""
    try:
        from docling.document_converter import DocumentConverter

        converter = DocumentConverter()
        result = converter.convert(file_path)
        doc = result.document

        # Export as markdown (preserves tables + structure)
        full_text = doc.export_to_markdown()
        total_pages = len(result.pages) if hasattr(result, 'pages') else 0

        # Extract global context from beginning
        global_context = _extract_global_context(full_text, carrier)

        # Chunk based on format config boundary patterns, or treat as single section
        if fmt_config and fmt_config.chunking.boundary_pattern:
            sections, _ = _chunk_by_boundary(full_text, global_context, fmt_config, file_path)
        else:
            sections = [ParsedSection(
                text=full_text,
                global_context=global_context,
                section_type="full_document",
            )]

        # Same safety net as parse_raw_text — cap any oversized section so an
        # unconfigured carrier doesn't ship a single 1.9M-token blob to Gemini.
        sections = _enforce_max_section_size(sections)

        return ParsedDocument(
            file_path=file_path,
            carrier=carrier,
            document_type=document_type,
            format_variant=fmt_config.name if fmt_config else "unknown",
            total_pages=total_pages,
            sections=sections,
        )

    except ImportError:
        logger.warning("Docling not available. Falling back to raw text.")
        return parse_raw_text(file_path, carrier, document_type, fmt_config)
    except Exception as e:
        logger.error(f"Docling parse failed for {file_path}: {e}. Falling back to raw text.")
        return parse_raw_text(file_path, carrier, document_type, fmt_config)


# ============================================
# Path 2: Raw Text (Mainframe CSRs, simple PDFs)
# ============================================

def parse_raw_text(
    file_path: str,
    carrier: str,
    document_type: str,
    fmt_config: FormatConfig | None = None,
) -> ParsedDocument:
    """Parse text-based PDFs with pdfplumber. For mainframe CSRs and simple layouts."""
    column_aware = fmt_config.chunking.column_aware if fmt_config else False
    full_text = _extract_raw_text(file_path, column_aware=column_aware)

    if not full_text or not full_text.strip():
        # No text extracted — likely scanned PDF. Flag for multimodal extraction.
        # For multi-page scanned PDFs, create one section per page group so each
        # gets its own Gemini multimodal call. A single call for a large scanned
        # PDF hits output token limits and only extracts the first few pages.
        total_pages = _count_pdf_pages(file_path)
        logger.info(f"No text in {Path(file_path).name}. Flagging {total_pages} pages "
                     f"as scanned for multimodal extraction.")

        # Group pages for multimodal extraction. Fewer pages per group = better OCR
        # for dense scanned documents, but more API calls. Default 3; configurable
        # via format config chunking.scanned_pages_per_group.
        pages_per_group = 3
        if fmt_config and hasattr(fmt_config.chunking, 'scanned_pages_per_group'):
            pages_per_group = fmt_config.chunking.scanned_pages_per_group or pages_per_group
        sections = []
        for start_page in range(1, total_pages + 1, pages_per_group):
            end_page = min(start_page + pages_per_group - 1, total_pages)
            page_range = f"{start_page}-{end_page}" if start_page != end_page else str(start_page)
            sections.append(ParsedSection(
                text=f"[SCANNED_PDF pages {page_range} — use multimodal extraction.]",
                section_type="scanned",
                global_context=f"Carrier: {carrier}, File: {Path(file_path).name}, Pages: {page_range}",
            ))

        return ParsedDocument(
            file_path=file_path,
            carrier=carrier,
            document_type=document_type,
            format_variant="scanned",
            total_pages=total_pages,
            sections=sections,
        )

    total_pages = _count_pdf_pages(file_path)
    global_context = _extract_global_context(full_text, carrier)
    validation_data = None

    # Enhance global context for AT&T box-format CSRs with address lookup table.
    # Box-format CSRs have SLA entries and LA/SA addresses scattered across pages
    # that won't appear in every 3-page chunk. Pre-extracting the address table
    # and injecting it ensures every chunk has the full address context.
    format_name = fmt_config.name if fmt_config else ""
    if "box" in format_name.lower() and document_type == "csr":
        addr_context = _extract_att_csr_address_context(full_text)
        if addr_context:
            global_context += addr_context

    # Chunk based on format config
    pages_per_chunk = getattr(fmt_config.chunking, 'pages_per_chunk', None) if fmt_config else None

    if pages_per_chunk and file_path.endswith(".pdf"):
        sections = _chunk_by_pages(file_path, global_context, pages_per_chunk)
    elif fmt_config and fmt_config.chunking.section_markers:
        sections = _chunk_by_section_markers(full_text, global_context, fmt_config)
        # Extract validation data if configured
        if fmt_config.chunking.validation_section:
            validation_data = _extract_validation_section(full_text, fmt_config.chunking.validation_section)
    elif fmt_config and fmt_config.chunking.boundary_pattern:
        sections, validation_data = _chunk_by_boundary(
            full_text, global_context, fmt_config, file_path
        )
    else:
        # No chunking config — treat entire document as one section
        sections = [ParsedSection(
            text=full_text,
            global_context=global_context,
            section_type="full_document",
        )]

    # Safety net: cap section text to MAX_SECTION_CHARS so an unconfigured
    # carrier (no boundary_pattern in carrier.yaml) doesn't ship a 1.9M-token
    # blob to Gemini and silently fail with 400 INVALID_ARGUMENT.
    sections = _enforce_max_section_size(sections)

    return ParsedDocument(
        file_path=file_path,
        carrier=carrier,
        document_type=document_type,
        format_variant=fmt_config.name if fmt_config else "unknown",
        total_pages=total_pages,
        sections=sections,
        validation_data=validation_data,
    )


# ============================================
# Path 3: Structured Data (CSV/XLSX)
# ============================================

def parse_structured_data(
    file_path: str,
    carrier: str,
    document_type: str,
    fmt_config: FormatConfig | None = None,
) -> ParsedDocument:
    """Parse CSV/XLSX with pandas. Minimal LLM needed."""
    suffix = Path(file_path).suffix.lower()

    try:
        if suffix == ".csv":
            df = pd.read_csv(file_path)
        elif suffix == ".tsv":
            df = pd.read_csv(file_path, sep="\t")
        elif suffix == ".xlsx":
            df = pd.read_excel(file_path, engine="openpyxl")
        elif suffix == ".xls":
            df = pd.read_excel(file_path, engine="xlrd")
        else:
            logger.warning(f"Unsupported structured format: {suffix}")
            return ParsedDocument(
                file_path=file_path, carrier=carrier,
                document_type=document_type, format_variant="unknown",
            )

        columns = list(df.columns)
        global_ctx = f"File: {Path(file_path).name}, Columns: {columns}, Total Rows: {len(df)}"

        # Chunk large dataframes into batches of 100 rows
        chunk_size = 100
        sections = []
        for start in range(0, len(df), chunk_size):
            chunk_df = df.iloc[start:start + chunk_size]
            text = f"Columns: {columns}\n\nData (rows {start+1}-{start+len(chunk_df)} of {len(df)}):\n{chunk_df.to_string()}"
            rows_as_strings = chunk_df.fillna("").astype(str).values.tolist()
            tables = [[columns] + rows_as_strings]

            sections.append(ParsedSection(
                text=text,
                tables=tables,
                section_type="structured_data",
                global_context=global_ctx,
            ))

        logger.info(f"Structured data: {len(df)} rows → {len(sections)} chunks of {chunk_size}")

        return ParsedDocument(
            file_path=file_path,
            carrier=carrier,
            document_type=document_type,
            format_variant=fmt_config.name if fmt_config else "structured",
            sections=sections,
        )

    except Exception as e:
        logger.error(f"Structured data parse failed for {file_path}: {e}")
        return ParsedDocument(
            file_path=file_path, carrier=carrier,
            document_type=document_type, format_variant="unknown",
        )


# ============================================
# Chunking Helpers
# ============================================

def _extract_raw_text(
    file_path: str,
    max_pages: int | None = None,
    column_aware: bool = False,
) -> str:
    """Extract all text from PDF using pdfplumber.

    If column_aware=True, detect two-column layouts and extract left column
    then right column separately (instead of interleaving them left-to-right).
    """
    suffix = Path(file_path).suffix.lower()
    if suffix != ".pdf":
        try:
            with open(file_path, "r", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""

    try:
        with pdfplumber.open(file_path) as pdf:
            pages = pdf.pages[:max_pages] if max_pages else pdf.pages
            texts = []
            for page in pages:
                if column_aware:
                    text = _extract_page_column_aware(page)
                else:
                    text = page.extract_text()
                if text:
                    texts.append(text)
            return "\n\n".join(texts)
    except Exception as e:
        logger.error(f"PDF text extraction failed: {e}")
        return ""


def _extract_page_column_aware(page) -> str:
    """Extract text from a page, separating left and right columns.

    For pages with two-column layouts (like Spectrum consolidated invoices),
    this prevents interleaving of column content. Words in the left column
    are extracted first, then words in the right column, preserving the
    spatial relationship between account numbers, addresses, and charges.
    """
    words = page.extract_words(keep_blank_chars=True, x_tolerance=3, y_tolerance=3)
    if not words:
        return page.extract_text() or ""

    page_midpoint = page.width / 2

    # Group words into lines by y-proximity
    words.sort(key=lambda w: (w['top'], w['x0']))
    lines = []
    current_line = []
    current_y = None

    for w in words:
        if current_y is None or abs(w['top'] - current_y) < 5:
            current_line.append(w)
            if current_y is None:
                current_y = w['top']
        else:
            if current_line:
                lines.append((current_y, current_line))
            current_line = [w]
            current_y = w['top']
    if current_line:
        lines.append((current_y, current_line))

    # Detect two-column regions: lines where words exist in both halves
    two_col_lines = 0
    for _, line_words in lines:
        left = any(w['x0'] < page_midpoint - 20 for w in line_words)
        right = any(w['x0'] > page_midpoint + 20 for w in line_words)
        if left and right:
            two_col_lines += 1

    # If less than 20% of lines are two-column, just use standard extraction
    if two_col_lines < len(lines) * 0.2:
        return page.extract_text() or ""

    # Extract text by column: full-width lines go inline,
    # two-column regions are split into left-then-right blocks
    result_parts = []
    left_buffer = []
    right_buffer = []
    in_two_col = False

    for _, line_words in lines:
        left_words = [w for w in line_words if w['x0'] < page_midpoint]
        right_words = [w for w in line_words if w['x0'] >= page_midpoint]

        has_left = bool(left_words)
        has_right = bool(right_words)

        if has_left and has_right:
            # Two-column line
            if not in_two_col:
                in_two_col = True
            left_text = " ".join(w['text'] for w in sorted(left_words, key=lambda w: w['x0']))
            right_text = " ".join(w['text'] for w in sorted(right_words, key=lambda w: w['x0']))
            left_buffer.append(left_text)
            right_buffer.append(right_text)
        else:
            # Full-width line — flush any buffered columns first
            if in_two_col and (left_buffer or right_buffer):
                result_parts.append("\n".join(left_buffer))
                result_parts.append("\n".join(right_buffer))
                left_buffer = []
                right_buffer = []
                in_two_col = False

            line_text = " ".join(w['text'] for w in sorted(line_words, key=lambda w: w['x0']))
            result_parts.append(line_text)

    # Flush remaining buffers
    if left_buffer or right_buffer:
        result_parts.append("\n".join(left_buffer))
        result_parts.append("\n".join(right_buffer))

    return "\n".join(result_parts)


def _count_pdf_pages(file_path: str) -> int:
    """Count pages in PDF."""
    if not file_path.endswith(".pdf"):
        return 0
    try:
        with pdfplumber.open(file_path) as pdf:
            return len(pdf.pages)
    except Exception:
        return 0


def _extract_global_context(text: str, carrier: str) -> str:
    """Extract account-level context from first ~500 chars of document."""
    header = text[:500]
    return f"Carrier: {carrier}\n---\n{header}"


def _extract_att_csr_address_context(full_text: str) -> str:
    """Extract address lookup table from AT&T box-format CSR.

    Box-format CSRs have address information in two places:
    1. MAIN LISTINGS section: LA/SA fields — the default address for all TNs
    2. LEFT HAND FIDS section: SLA entries — per-station addresses

    Per-TN USOC lines reference their station via /SLA NNN. This function
    builds the lookup table so the LLM can assign the correct address to
    each TN during extraction.

    Returns a string to append to global_context, or empty string if no
    address information found.
    """
    lines = []

    # 1. Extract main LA/SA address (default for TNs without SLA)
    la_match = re.search(r'\bLA\s+[!|]([^!\n]+)', full_text)
    sa_match = re.search(r'\bSA\s+[!|]([^!\n]+)', full_text)
    dzip_match = re.search(r'\bDZIP\s+[!|](\d+)', full_text)

    default_addr = (sa_match or la_match)
    if default_addr:
        addr = default_addr.group(1).strip().rstrip(',')
        city = ""
        # LA/SA format: "ADDRESS, CITY" — extract city from after the comma
        if ',' in addr:
            parts = addr.rsplit(',', 1)
            addr = parts[0].strip()
            city = parts[1].strip()
        zip_code = dzip_match.group(1).strip() if dzip_match else ""
        lines.append(f"DEFAULT_ADDRESS: {addr}")
        if city:
            lines.append(f"DEFAULT_CITY: {city}")
        if zip_code:
            lines.append(f"DEFAULT_ZIP: {zip_code}")

    # 2. Extract SLA lookup table from LEFT HAND FIDS
    sla_entries = re.findall(r'\bSLA\s+[!|](\d+)[-–]([^!\n]+)', full_text)
    if sla_entries:
        lines.append("SLA_LOOKUP_TABLE:")
        for sla_num, sla_addr in sla_entries:
            addr = sla_addr.strip().rstrip(',')
            # Strip /LSO, /DPI, /SN suffixes (internal codes, not address)
            addr = re.sub(r'\s*/(?:LSO|DPI|SN)\b.*$', '', addr).strip()
            # Split "ADDRESS, CITY" if present
            if ',' in addr:
                parts = addr.rsplit(',', 1)
                addr_part = parts[0].strip()
                city_part = parts[1].strip()
                lines.append(f"  SLA {sla_num} = {addr_part}, {city_part}")
            else:
                lines.append(f"  SLA {sla_num} = {addr}")

    # 3. Extract contract info (/CNUM, /TA) from global CSR context.
    # /CNUM appears in equipment sections — it's the contract reference number.
    # /TA appears on per-TN USOC lines — "term_months, date".
    # Injecting these into global_context ensures every chunk has contract info
    # even if /CNUM or /TA only appears on one page.
    cnum_matches = re.findall(r'/CNUM\s+(\S+)', full_text)
    if cnum_matches:
        # Deduplicate — typically the same /CNUM appears on multiple lines
        unique_cnum = sorted(set(cnum_matches))
        lines.append(f"CONTRACT_NUMBER: {unique_cnum[0]}")
        if len(unique_cnum) > 1:
            lines.append(f"CONTRACT_NUMBERS_ALT: {', '.join(unique_cnum[1:])}")

    if not lines:
        return ""

    context = "\n--- ADDRESS LOOKUP TABLE ---\n"
    context += "\n".join(lines)
    context += "\n--- END ADDRESS TABLE ---\n"

    logger.info(f"AT&T CSR address context: default={bool(default_addr)}, "
                f"{len(sla_entries)} SLA entries, {len(cnum_matches)} /CNUM refs")
    return context


def _enforce_max_section_size(
    sections: list[ParsedSection],
    max_chars: int = MAX_SECTION_CHARS,
) -> list[ParsedSection]:
    """Safety net: split any oversized section so no single LLM call blows
    past the input-token limit.

    Without this, an unconfigured carrier (e.g. Verizon — no boundary pattern
    in carrier.yaml) produces one "full_document" section. For a 9 MB PDF that
    can be ~1.9M tokens and Gemini Flash silently rejects it with 400
    INVALID_ARGUMENT. The user sees 0 rows and no error — the original
    "extraction not processing properly" symptom.

    Splits on `\n\n` (paragraph), then `\n` (line) boundaries. If a single
    line is itself > max_chars (rare; pathological PDF text) it's hard-split.
    Each piece carries the same global_context, sub_account, and section_type
    so downstream extraction logic still groups them as one logical section.
    """
    out: list[ParsedSection] = []
    for section in sections:
        text = section.text or ""
        if len(text) <= max_chars:
            out.append(section)
            continue

        pieces = _split_text_safely(text, max_chars)
        logger.warning(
            "Section text %d chars > %d cap — splitting into %d pieces "
            "(carrier=%s, sub_account=%s, section_type=%s)",
            len(text), max_chars, len(pieces),
            getattr(section, "carrier", None), section.sub_account, section.section_type,
        )
        for i, piece in enumerate(pieces, 1):
            out.append(ParsedSection(
                text=piece,
                global_context=section.global_context,
                section_type=section.section_type,
                sub_account=section.sub_account,
            ))
    return out


def _split_text_safely(text: str, max_chars: int) -> list[str]:
    """Split text into ≤max_chars chunks, preferring paragraph then line
    boundaries. Last-resort hard-split for pathological cases."""
    if len(text) <= max_chars:
        return [text]

    # Try paragraph-aligned split first.
    chunks = _greedy_pack(text.split("\n\n"), separator="\n\n", max_chars=max_chars)
    if all(len(c) <= max_chars for c in chunks):
        return chunks

    # Fall back to line-aligned within any chunk that's still too large.
    final: list[str] = []
    for c in chunks:
        if len(c) <= max_chars:
            final.append(c)
            continue
        line_chunks = _greedy_pack(c.split("\n"), separator="\n", max_chars=max_chars)
        for lc in line_chunks:
            if len(lc) <= max_chars:
                final.append(lc)
            else:
                # Truly pathological — hard-split.
                for i in range(0, len(lc), max_chars):
                    final.append(lc[i:i + max_chars])
    return final


def _greedy_pack(parts: list[str], separator: str, max_chars: int) -> list[str]:
    """Pack `parts` back into chunks ≤max_chars, joined with `separator`.
    Each individual part may still exceed max_chars — caller handles."""
    chunks: list[str] = []
    cur: list[str] = []
    cur_len = 0
    sep_len = len(separator)
    for p in parts:
        added = len(p) + (sep_len if cur else 0)
        if cur and cur_len + added > max_chars:
            chunks.append(separator.join(cur))
            cur = [p]
            cur_len = len(p)
        else:
            cur.append(p)
            cur_len += added
    if cur:
        chunks.append(separator.join(cur))
    return chunks


def _chunk_by_boundary(
    full_text: str,
    global_context: str,
    fmt_config: FormatConfig,
    file_path: str,
) -> tuple[list[ParsedSection], dict | None]:
    """Split document at boundary pattern matches.

    Used for Windstream Enterprise (ACTIVITY FOR ACCOUNT) and AT&T invoices (Chargesfor).
    """
    pattern = fmt_config.chunking.boundary_pattern
    validation_data = None

    # Extract validation section if configured
    if fmt_config.chunking.validation_section:
        validation_data = _extract_validation_section(full_text, fmt_config.chunking.validation_section)

    # Find all boundary positions
    matches = list(re.finditer(pattern, full_text))

    if not matches:
        # No boundaries found — return whole doc as one section
        return [ParsedSection(
            text=full_text,
            global_context=global_context,
            section_type="full_document",
        )], validation_data

    sections = []

    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        chunk_text = full_text[start:end].strip()

        if not chunk_text:
            continue

        # Try to extract sub-account from boundary match
        sub_account = match.group(1) if match.lastindex else None

        sections.append(ParsedSection(
            text=chunk_text,
            global_context=global_context,
            section_type="sub_account",
            sub_account=sub_account,
        ))

    logger.info(f"Chunked {Path(file_path).name} into {len(sections)} sections by boundary pattern")
    return sections, validation_data


def _extract_address_blocks(full_text: str, boundary_matches: list) -> dict[str, str]:
    """Extract address blocks for each account using line-based scanning.

    Falls back to basic scanning when spatial extraction is not available.
    """
    # This is the fallback — spatial extraction is preferred (see _extract_address_blocks_spatial)
    result = {}
    lines = full_text.split("\n")

    for match in boundary_matches:
        account = match.group(1) if match.lastindex else None
        if not account:
            continue

        text_before = full_text[:match.start()]
        line_num = text_before.count("\n")

        address_lines = []
        stop_patterns = re.compile(
            r'(AccountNumber:|TotalForAccount|MonthlyCharges|SecurityCode:|'
            r'ChargeDetails|Date\s+Description|Subtotal|PastDue)',
            re.IGNORECASE
        )

        for j in range(line_num - 1, max(line_num - 6, -1), -1):
            if j < 0 or j >= len(lines):
                break
            line = lines[j].strip()
            if not line:
                break
            if stop_patterns.search(line):
                break
            address_lines.insert(0, line)

        if address_lines:
            result[account] = "\n".join(address_lines)

    return result


def _extract_address_blocks_spatial(file_path: str) -> dict[str, str]:
    """Extract address blocks using PDF spatial coordinates.

    For each AccountNumber: word on each page, finds the words directly above it
    in the same column (similar x-position) to build the address block.
    This is immune to column-interleaving issues.
    """
    result = {}
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                words = page.extract_words(keep_blank_chars=True, x_tolerance=3, y_tolerance=3)
                if not words:
                    continue

                # Find all AccountNumber: words
                for w in words:
                    text = w['text'].replace(' ', '')
                    if not text.startswith('AccountNumber:'):
                        continue

                    account_match = re.search(r'(\d{13,16})', text)
                    if not account_match:
                        continue
                    account = account_match.group(1)

                    # Get the x-position and y-position of this AccountNumber word
                    acct_x0 = w['x0']
                    acct_y = w['top']

                    # Find words in the same column (x within 50px) and directly above.
                    # Address blocks are compact: name + street + city/state/zip = ~40-50px tall.
                    same_col_above = [
                        ow for ow in words
                        if abs(ow['x0'] - acct_x0) < 50
                        and ow['top'] < acct_y
                        and acct_y - ow['top'] < 60  # tight: only the address block
                    ]
                    same_col_above.sort(key=lambda x: x['top'], reverse=True)

                    # Group into lines by y-proximity, stopping at charge/metadata lines
                    address_lines = []
                    stop_re = re.compile(
                        r'(AccountNumber|TotalForAccount|SecurityCode|MonthlyCharges|'
                        r'Subtotal|PastDue|Adjustments|ChargeDetails|^\$|'
                        r'^\d+\.\d{2}$|^[A-Z][a-z]{2}\d{2})'  # dollar amounts, date prefixes
                    )
                    for ow in same_col_above:
                        ow_text = ow['text'].strip()
                        if not ow_text:
                            continue
                        if stop_re.search(ow_text):
                            break
                        # Check if this is on the same line as a previous word
                        if address_lines and abs(ow['top'] - address_lines[-1][1]) < 5:
                            address_lines[-1] = (address_lines[-1][0] + " " + ow_text, ow['top'])
                        else:
                            address_lines.append((ow_text, ow['top']))

                    if address_lines:
                        # Reverse to get top-to-bottom order
                        address_lines.reverse()

                        # Also grab the zip code — it's often a separate word on the
                        # same line as CITY,STATE but at a different x-position
                        zip_match = None
                        for ow in words:
                            if (abs(ow['top'] - address_lines[-1][1]) < 5  # same y as last address line
                                and ow['x0'] > acct_x0  # to the right
                                and re.match(r'^\d{5}(-\d{4})?$', ow['text'].strip())):
                                zip_match = ow['text'].strip()
                                break

                        addr_text = "\n".join(line[0] for line in address_lines)
                        if zip_match:
                            addr_text += f" {zip_match}"
                        result[account] = addr_text

    except Exception as e:
        logger.warning(f"Spatial address extraction failed: {e}")

    return result


def _chunk_by_section_markers(
    full_text: str,
    global_context: str,
    fmt_config: FormatConfig,
) -> list[ParsedSection]:
    """Split document at section marker lines (e.g., ---LISTINGS---, ---BILL---).

    Used for AT&T CSR section-marker format.
    """
    markers = fmt_config.chunking.section_markers
    sections = []

    # Find positions of each marker
    marker_positions = []
    for marker in markers:
        for match in re.finditer(re.escape(marker), full_text):
            marker_positions.append((match.start(), marker))

    marker_positions.sort(key=lambda x: x[0])

    if not marker_positions:
        return [ParsedSection(
            text=full_text,
            global_context=global_context,
            section_type="full_document",
        )]

    for i, (pos, marker) in enumerate(marker_positions):
        end = marker_positions[i + 1][0] if i + 1 < len(marker_positions) else len(full_text)
        chunk_text = full_text[pos:end].strip()

        if chunk_text:
            section_name = marker.replace("---", "").strip().lower()
            sections.append(ParsedSection(
                text=chunk_text,
                global_context=global_context,
                section_type=section_name,
            ))

    return sections


def _chunk_by_pages(
    file_path: str,
    global_context: str,
    pages_per_chunk: int,
) -> list[ParsedSection]:
    """Split PDF into chunks of N pages each."""
    sections = []
    try:
        with pdfplumber.open(file_path) as pdf:
            total = len(pdf.pages)
            for start in range(0, total, pages_per_chunk):
                end = min(start + pages_per_chunk, total)
                texts = []
                for i in range(start, end):
                    text = pdf.pages[i].extract_text()
                    if text:
                        texts.append(text)
                chunk_text = "\n\n".join(texts)
                if chunk_text.strip():
                    sections.append(ParsedSection(
                        text=chunk_text,
                        page_numbers=list(range(start + 1, end + 1)),
                        global_context=global_context,
                        section_type=f"pages_{start+1}-{end}",
                    ))
            logger.info(f"Page-chunked {Path(file_path).name}: {total} pages → {len(sections)} chunks of {pages_per_chunk}")
    except Exception as e:
        logger.error(f"Page chunking failed: {e}")
    return sections


def _extract_validation_section(full_text: str, section_name: str) -> dict | None:
    """Extract validation data (e.g., Windstream LOCATION SUMMARY totals)."""
    if section_name == "LOCATION SUMMARY":
        return _parse_windstream_location_summary(full_text)
    elif section_name == "---REVENUE AMOUNTS---":
        return _parse_att_revenue_amounts(full_text)
    elif section_name == "SERV & EQUIP ACCOUNT SUMMARY":
        return _parse_att_account_summary(full_text)
    return None


def _parse_windstream_location_summary(text: str) -> dict | None:
    """Parse LOCATION SUMMARY table from Windstream Enterprise invoice.

    Returns dict mapping sub-account numbers to their total charges.
    """
    match = re.search(r"LOCATION SUMMARY.*?\n(.*?)(?=\n\n|\nACTIVITY FOR)", text, re.DOTALL)
    if not match:
        return None

    summary = {}
    for line in match.group(1).split("\n"):
        # Pattern: account_number  account_name  monthly  usage  other  surcharges  total
        parts = line.strip().split()
        if parts and re.match(r"\d{5,10}", parts[0]):
            account = parts[0]
            # Find dollar amounts (last value is TOTAL)
            amounts = re.findall(r"\$[\d,]+\.\d{2}", line)
            if amounts:
                total = amounts[-1].replace("$", "").replace(",", "")
                summary[account] = float(total)

    return {"location_summary": summary} if summary else None


def _parse_att_revenue_amounts(text: str) -> dict | None:
    """Parse ---REVENUE AMOUNTS--- section from AT&T CSR."""
    match = re.search(r"---REVENUE AMOUNTS---\s*(.*?)(?=---|\Z)", text, re.DOTALL)
    if not match:
        return None

    total_match = re.search(r"LOCAL SERVICE TOTAL\s+([\d.]+)", match.group(1))
    if total_match:
        return {"revenue_total": float(total_match.group(1))}
    return None


def _parse_att_account_summary(text: str) -> dict | None:
    """Parse SERV & EQUIP ACCOUNT SUMMARY from AT&T CSR box format.

    Returns USOC → description mapping.
    """
    match = re.search(r"SERV & EQUIP ACCOUNT SUMMARY(.*?)(?=\!-{10,}|\Z)", text, re.DOTALL)
    if not match:
        return None

    usoc_map = {}
    for line in match.group(1).split("\n"):
        # Pattern: ! ! 1 !CODE Description !
        code_match = re.search(r"!\s*\d+\s*!(\w+)\s+(.+?)!", line)
        if code_match:
            usoc_map[code_match.group(1).strip()] = code_match.group(2).strip()

    return {"usoc_summary": usoc_map} if usoc_map else None


# ============================================
# Path 4: Email (.msg, .eml)
# ============================================

def parse_email(
    file_path: str,
    carrier: str,
    document_type: str,
) -> ParsedDocument:
    """Parse email files (.msg, .eml) into sections for LLM extraction."""
    from backend.pipeline.classifier import _extract_email_text

    text = _extract_email_text(file_path)
    if not text:
        logger.warning(f"No text extracted from email: {file_path}")
        return ParsedDocument(
            file_path=file_path, carrier=carrier,
            document_type=document_type or "email",
            format_variant="email",
        )

    sections = [ParsedSection(
        text=text,
        section_type="email_body",
        page_numbers=[1],
    )]

    # Check for PDF attachments in .msg files
    suffix = Path(file_path).suffix.lower()
    if suffix == ".msg":
        try:
            import extract_msg
            msg = extract_msg.Message(file_path)
            for att in msg.attachments:
                if hasattr(att, 'filename') and att.filename and att.filename.lower().endswith('.pdf'):
                    logger.info(f"Email has PDF attachment: {att.filename}")
                    # Note: attachment extraction would save to temp and process separately
                    # For now, just log it — full attachment handling is a follow-up
        except Exception as e:
            logger.debug(f"Attachment check failed: {e}")

    return ParsedDocument(
        file_path=file_path,
        carrier=carrier,
        document_type=document_type or "email",
        format_variant="email",
        total_pages=1,
        sections=sections,
    )


# ============================================
# Path 5: Word Documents (.docx)
# ============================================

def parse_docx(
    file_path: str,
    carrier: str,
    document_type: str,
) -> ParsedDocument:
    """Parse Word documents into sections for LLM extraction."""
    try:
        import docx
        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        # Also extract tables
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_texts.append("\n".join(rows))

        if table_texts:
            text += "\n\nTABLES:\n" + "\n\n".join(table_texts)

    except Exception as e:
        logger.warning(f"DOCX parse failed: {e}")
        text = ""

    if not text:
        return ParsedDocument(
            file_path=file_path, carrier=carrier,
            document_type=document_type or "report",
            format_variant="docx",
        )

    sections = [ParsedSection(
        text=text,
        section_type="document",
        page_numbers=[1],
    )]

    return ParsedDocument(
        file_path=file_path,
        carrier=carrier,
        document_type=document_type or "report",
        format_variant="docx",
        total_pages=1,
        sections=sections,
    )
